import asyncio
import io
import logging
import os
import uuid
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import pdf2image  # Add pdf2image for PDF to image conversion
import pymongo
from PIL import Image
from pymilvus import Collection, connections
from tqdm import tqdm

from app.core.config import settings
from app.db.minio import async_minio_manager
from app.rag.colbert_service import colbert, logger

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize MongoDB connection
try:
    mongo_client = pymongo.MongoClient(
        settings.mongodb_uri, serverSelectionTimeoutMS=2000
    )
    mongo_client.server_info()  # Will raise an exception if cannot connect
    mongo_db = mongo_client[settings.mongodb_db]
    docs_collection = mongo_db["documents"]
    chunks_collection = mongo_db["document_chunks"]
    logger.info("Document processor connected to MongoDB")
except Exception as e:
    logger.warning(f"Document processor could not connect to MongoDB: {e}")
    mongo_client = None
    mongo_db = None
    docs_collection = None
    chunks_collection = None

# Initialize Milvus connection
try:
    connections.connect("default", host=settings.milvus_host, port=settings.milvus_port)
    logger.info("Document processor connected to Milvus")
except Exception as e:
    logger.warning(f"Document processor could not connect to Milvus: {e}")


class DocumentProcessor:
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        # Create a directory for PDF page images
        self.pdf_images_dir = self.upload_dir / "pdf_images"
        self.pdf_images_dir.mkdir(exist_ok=True)

    def process_file(
        self,
        file_content: bytes,
        filename: str,
        kb_id: str,
        chunk_size: int = None,
        chunk_overlap: int = None,
    ) -> str:
        """Process a file and store its chunks in MongoDB and Milvus"""
        # Generate a unique document ID
        doc_id = str(uuid.uuid4())
        logger.info(f"Processing file: {filename} (ID: {doc_id})")

        # Save file to disk
        file_path = self.upload_dir / filename
        with open(file_path, "wb") as f:
            f.write(file_content)
        logger.info(f"Saved file to: {file_path}")

        # Extract text and metadata based on file type
        text, metadata = self._extract_text_and_metadata(file_path, filename)
        logger.info(f"Extracted metadata: {metadata}")

        # Store document metadata
        if docs_collection:
            try:
                docs_collection.insert_one(
                    {
                        "_id": doc_id,
                        "filename": filename,
                        "kb_id": kb_id,
                        "metadata": metadata,
                        "file_path": str(file_path),
                    }
                )
                logger.info(f"Stored document metadata in MongoDB")
            except Exception as e:
                logger.error(f"Error storing document metadata: {e}")
        else:
            logger.warning("MongoDB not available, skipping document metadata storage")

        # For PDFs, process each page as an image
        if file_path.suffix.lower() == ".pdf":
            logger.info(f"Processing PDF file: {filename}")
            self._process_pdf_as_images(file_path, doc_id, kb_id)
        else:
            # For non-PDF files, continue with text chunking
            # Split text into chunks
            chunk_size = chunk_size or settings.chunk_size
            chunk_overlap = chunk_overlap or settings.chunk_overlap
            chunks = self._split_text(text, chunk_size, chunk_overlap)
            logger.info(f"Split text into {len(chunks)} chunks")

            # Process chunks
            self._process_chunks(chunks, doc_id, kb_id)

        return doc_id

    def _process_pdf_as_images(self, pdf_path: Path, doc_id: str, kb_id: str) -> None:
        """Process a PDF file by converting each page to an image and embedding it"""
        try:
            logger.info(f"Converting PDF to images: {pdf_path}")

            # Convert PDF to list of PIL Image objects
            pdf_images = pdf2image.convert_from_path(
                pdf_path,
                dpi=200,  # Reasonable resolution for document processing
                fmt="png",
            )

            logger.info(f"Converted PDF to {len(pdf_images)} pages/images")

            # Create a subdirectory for this document's images
            doc_image_dir = self.pdf_images_dir / doc_id
            doc_image_dir.mkdir(exist_ok=True)

            # Process each page image
            for i, image in enumerate(pdf_images):
                page_num = i + 1
                image_filename = f"{doc_id}_page_{page_num}.png"
                image_path = doc_image_dir / image_filename

                # Save image to disk
                image.save(image_path)
                logger.info(f"Saved PDF page {page_num} as image: {image_path}")

                # Create page metadata
                page_metadata = {
                    "source_file": pdf_path.name,
                    "page_number": page_num,
                    "total_pages": len(pdf_images),
                    "width": image.width,
                    "height": image.height,
                    "type": "pdf_page",
                    "format": "png",
                }

                # Process image embedding
                try:
                    # Generate embedding using ColBERT model
                    embeddings = colbert.process_image([image])

                    # Ensure Milvus collection exists
                    colbert.create_kb_collection(kb_id)

                    # Store embedding in Milvus
                    collection_name = f"colqwen_{kb_id}"
                    collection = Collection(collection_name)

                    # Insert embedding with unique chunk ID
                    chunk_id = str(uuid.uuid4())

                    entities = [
                        [chunk_id],  # chunk_id
                        [doc_id],  # doc_id
                        embeddings,  # embedding
                    ]
                    collection.insert(entities)
                    logger.info(f"Stored PDF page {page_num} embedding in Milvus")

                    # Store page metadata in MongoDB
                    if chunks_collection:
                        chunks_collection.insert_one(
                            {
                                "_id": chunk_id,
                                "doc_id": doc_id,
                                "kb_id": kb_id,
                                "type": "pdf_page",
                                "page_number": page_num,
                                "image_path": str(image_path),
                                "metadata": page_metadata,
                            }
                        )
                        logger.info(f"Stored PDF page {page_num} metadata in MongoDB")

                except Exception as e:
                    logger.error(f"Error processing PDF page {page_num}: {e}")

        except Exception as e:
            logger.error(f"Error converting PDF to images: {e}")

    def process_image(
        self,
        image: Union[bytes, Path, str],
        filename: str,
        kb_id: str,
    ) -> str:
        """Process an image and store its embedding in Milvus"""
        # Generate a unique document ID
        doc_id = str(uuid.uuid4())
        logger.info(f"Processing image: {filename} (ID: {doc_id})")

        # Convert to PIL Image if needed
        if isinstance(image, bytes):
            pil_image = Image.open(BytesIO(image))
            # Save image to disk
            file_path = self.upload_dir / filename
            pil_image.save(file_path)
            logger.info(f"Saved image from bytes to: {file_path}")
        elif isinstance(image, (str, Path)):
            file_path = Path(image)
            pil_image = Image.open(file_path)
            logger.info(f"Opened image from path: {file_path}")
        else:
            pil_image = image
            file_path = self.upload_dir / filename
            pil_image.save(file_path)
            logger.info(f"Saved image object to: {file_path}")

        # Extract metadata
        metadata = {
            "width": pil_image.width,
            "height": pil_image.height,
            "format": pil_image.format,
            "mode": pil_image.mode,
        }
        logger.info(f"Image metadata: {metadata}")

        # Store document metadata
        if docs_collection:
            try:
                docs_collection.insert_one(
                    {
                        "_id": doc_id,
                        "filename": filename,
                        "kb_id": kb_id,
                        "type": "image",
                        "metadata": metadata,
                        "file_path": str(file_path),
                    }
                )
                logger.info(f"Stored image metadata in MongoDB")
            except Exception as e:
                logger.error(f"Error storing image metadata: {e}")
        else:
            logger.warning("MongoDB not available, skipping image metadata storage")

        # Generate embedding
        try:
            embeddings = colbert.process_image([pil_image])
            logger.info(
                f"Generated image embedding with dimension: {len(embeddings[0])}"
            )

            # Ensure Milvus collection exists
            colbert.create_kb_collection(kb_id)

            # Store embedding in Milvus
            try:
                collection_name = f"colqwen_{kb_id}"
                collection = Collection(collection_name)

                # Insert embedding
                chunk_id = str(uuid.uuid4())
                entities = [
                    [chunk_id],  # chunk_id
                    [doc_id],  # doc_id
                    embeddings,  # embedding
                ]
                collection.insert(entities)
                logger.info(f"Stored image embedding in Milvus")

                # Store chunk metadata in MongoDB
                if chunks_collection:
                    try:
                        chunks_collection.insert_one(
                            {
                                "_id": chunk_id,
                                "doc_id": doc_id,
                                "kb_id": kb_id,
                                "type": "image",
                                "metadata": metadata,
                            }
                        )
                        logger.info(f"Stored chunk metadata in MongoDB")
                    except Exception as e:
                        logger.error(f"Error storing chunk metadata: {e}")
                else:
                    logger.warning(
                        "MongoDB not available, skipping chunk metadata storage"
                    )
            except Exception as e:
                logger.error(f"Error storing embedding in Milvus: {e}")
        except Exception as e:
            logger.error(f"Error generating image embedding: {e}")

        return doc_id

    def _extract_text_and_metadata(
        self, file_path: Path, filename: str
    ) -> Tuple[str, Dict]:
        """Extract text and metadata from a file based on its extension"""
        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return self._extract_from_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            return self._extract_from_docx(file_path)
        elif extension in [".txt", ".md", ".py", ".js", ".html", ".css"]:
            return self._extract_from_text(file_path)
        elif extension in [".jpg", ".jpeg", ".png", ".gif", ".bmp"]:
            # For images, we don't extract text but return empty text and image metadata
            img = Image.open(file_path)
            metadata = {
                "width": img.width,
                "height": img.height,
                "format": img.format,
                "mode": img.mode,
            }
            return "", metadata
        else:
            # For unsupported files, return empty text and basic metadata
            return "", {"filename": filename, "extension": extension}

    def _extract_from_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text and metadata from a PDF file"""
        try:
            # Import PyPDF2 here to avoid dependency issues if not installed
            import PyPDF2

            # Open the PDF file
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                # Extract text from all pages
                text = ""
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"

                # Extract metadata
                info = pdf_reader.metadata
                metadata = {
                    "type": "pdf",
                    "pages": num_pages,
                    "title": info.title if info.title else None,
                    "author": info.author if info.author else None,
                    "subject": info.subject if info.subject else None,
                    "creator": info.creator if info.creator else None,
                    "producer": info.producer if info.producer else None,
                }

                return text, metadata

        except ImportError:
            logger.warning("PyPDF2 not installed. Using basic PDF metadata extraction.")
            # Use pdf2image to count pages
            try:
                images = pdf2image.convert_from_path(
                    file_path, dpi=72, first_page=1, last_page=1
                )
                return "", {"type": "pdf", "pages": len(images)}
            except Exception as e:
                logger.error(f"Error counting PDF pages: {e}")
                return "", {"type": "pdf", "pages": 1}
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return "", {"error": str(e)}

    def _extract_from_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text and metadata from a DOCX file"""
        try:
            # Import python-docx here to avoid dependency issues if not installed
            import docx

            doc = docx.Document(file_path)

            # Extract text from paragraphs
            text = "\n".join([para.text for para in doc.paragraphs])

            # Extract basic metadata
            metadata = {
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections),
            }

            return text, metadata

        except ImportError:
            logger.warning("python-docx not installed. Using placeholder text.")
            return f"Text extracted from DOCX: {file_path.name}", {
                "type": "docx",
                "pages": 1,
            }
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return "", {"error": str(e)}

    def _extract_from_text(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text from a plain text file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            return text, {"type": "text", "extension": file_path.suffix}
        except Exception as e:
            logger.error(f"Error extracting text from text file: {e}")
            return "", {"error": str(e)}

    def _split_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """Split text into chunks of specified size with overlap"""
        if not text:
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + chunk_size, text_len)
            chunks.append(text[start:end])
            start += chunk_size - chunk_overlap

        return chunks

    def _process_chunks(self, chunks: List[str], doc_id: str, kb_id: str) -> None:
        """Process text chunks and store them in MongoDB and Milvus"""
        if not chunks:
            logger.info("No chunks to process")
            return

        logger.info(f"Processing {len(chunks)} chunks for document {doc_id}")

        try:
            # Generate embeddings for all chunks
            embeddings = colbert.process_query(chunks)
            logger.info(f"Generated embeddings for {len(chunks)} chunks")

            # Ensure Milvus collection exists
            colbert.create_kb_collection(kb_id)

            try:
                # Get collection
                collection_name = f"colqwen_{kb_id}"
                collection = Collection(collection_name)

                # Process chunks
                chunk_ids = []
                doc_ids = []

                # Store chunks in MongoDB and prepare for Milvus
                for i, chunk in enumerate(chunks):
                    chunk_id = str(uuid.uuid4())
                    chunk_ids.append(chunk_id)
                    doc_ids.append(doc_id)

                    # Store in MongoDB
                    if chunks_collection:
                        try:
                            chunks_collection.insert_one(
                                {
                                    "_id": chunk_id,
                                    "doc_id": doc_id,
                                    "kb_id": kb_id,
                                    "text": chunk,
                                    "chunk_index": i,
                                    "type": "text",
                                }
                            )
                        except Exception as e:
                            logger.error(f"Error storing chunk in MongoDB: {e}")
                    else:
                        logger.warning("MongoDB not available, skipping chunk storage")

                # Insert all embeddings at once
                if chunk_ids:
                    try:
                        entities = [
                            chunk_ids,  # chunk_id
                            doc_ids,  # doc_id
                            embeddings,  # embedding
                        ]
                        collection.insert(entities)
                        logger.info(f"Stored {len(chunks)} embeddings in Milvus")
                    except Exception as e:
                        logger.error(f"Error storing embeddings in Milvus: {e}")
            except Exception as e:
                logger.error(f"Error preparing Milvus collection: {e}")
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")


# Initialize the document processor
document_processor = DocumentProcessor()


async def convert_file_to_images(file_content: bytes):
    """
    Convert a file (PDF) to a list of image buffers
    Returns list of BytesIO objects containing images
    """
    # Create a temporary file for PDF processing
    with io.BytesIO(file_content) as pdf_buffer:
        try:
            # Convert PDF pages to list of PIL images
            images = pdf2image.convert_from_bytes(
                pdf_buffer.read(),
                dpi=200,
                fmt="png",
            )
            logger.info(f"Converted PDF to {len(images)} pages/images")

            # Convert PIL images to BytesIO buffers
            image_buffers = []
            for img in images:
                buffer = io.BytesIO()
                img.save(buffer, format="PNG")
                buffer.seek(0)
                image_buffers.append(buffer)

            return image_buffers
        except Exception as e:
            logger.error(f"Error converting PDF to images: {e}")
            raise


async def save_image_to_minio(
    username: str, original_filename: str, image_buffer: io.BytesIO
):
    """
    Save an image buffer to MinIO and return the filename and URL
    """
    try:
        # Generate a unique filename
        base_name = Path(original_filename).stem
        ext = ".png"
        unique_name = f"{base_name}_{uuid.uuid4()}{ext}"

        # Define the user's folder path
        folder = f"users/{username}/images"
        filename = f"{folder}/{unique_name}"

        # Upload to MinIO
        await async_minio_manager.upload_image(filename, image_buffer)

        # Generate a URL
        url = await async_minio_manager.create_presigned_url(filename)

        return filename, url
    except Exception as e:
        logger.error(f"Error saving image to MinIO: {e}")
        raise
